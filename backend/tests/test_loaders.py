"""
Loader tests: dispatch, plus one section per format.

The point of every test here is the same claim - a loader's job is to reach the
*shared* representation, so that whatever the source format was, the rest of
the pipeline sees blocks in reading order with headings marked and section
paths filled in. Each format then gets tests for the specific thing it has to
get right, which is different every time: code fences in Markdown, boilerplate
markup in HTML, style names in Word, header repetition in CSV.
"""
import io

import docx
import pytest

from rag.loaders import SUPPORTED_EXTENSIONS, load_document
from rag.loaders.csv import ROWS_PER_GROUP, parse_csv
from rag.loaders.html import parse_html
from rag.loaders.markdown import parse_markdown


def _kinds(blocks) -> list[str]:
    return [block.kind for block in blocks]


def _build_docx(build) -> bytes:
    document = docx.Document()
    build(document)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- Dispatch ---------------------------------------------------------------


def test_every_supported_extension_loads_something():
    samples = {
        ".md": b"# Title\n\nBody.",
        ".markdown": b"# Title\n\nBody.",
        ".txt": b"Just some text.",
        ".html": b"<h1>Title</h1><p>Body.</p>",
        ".htm": b"<h1>Title</h1><p>Body.</p>",
        ".csv": b"a,b\n1,2\n",
        ".tsv": b"a\tb\n1\t2\n",
        ".docx": _build_docx(lambda d: d.add_paragraph("Body.")),
    }
    # .pdf is covered end to end in test_pdf_loader.py; everything else has to
    # be represented here, so a new extension can't be added without a sample.
    assert set(samples) | {".pdf"} == set(SUPPORTED_EXTENSIONS)

    for extension, data in samples.items():
        pages = load_document(f"sample{extension}", data)
        assert len(pages) == 1, extension
        assert pages[0].blocks, extension


def test_unsupported_extension_is_rejected():
    with pytest.raises(ValueError, match="Unsupported file type"):
        load_document("photo.png", b"\x89PNG")


def test_empty_file_is_rejected():
    with pytest.raises(ValueError, match="empty"):
        load_document("notes.md", b"")


def test_extension_matching_ignores_case():
    assert load_document("NOTES.MD", b"# Title")[0].blocks


def test_pageless_formats_report_a_single_page():
    """
    A .md/.docx/.csv has no pages, so inventing them would put a number in
    every citation that corresponds to nothing a reader could look up.
    """
    pages = load_document("notes.md", b"# A\n\ntext\n\n# B\n\nmore text\n")

    assert [page.page_number for page in pages] == [1]


def test_loaders_produce_section_paths():
    """The shared contract: whatever the format, blocks know their section."""
    pages = load_document("notes.md", b"# Methods\n\nWe collected samples.\n")

    assert pages[0].blocks[1].section_path == ("Methods",)


# --- Markdown ---------------------------------------------------------------


def test_markdown_hash_headings_carry_their_own_level():
    blocks = parse_markdown("# One\n\ntext\n\n## Two\n\nmore\n")

    assert _kinds(blocks) == ["heading", "text", "heading", "text"]
    assert [block.level for block in blocks if block.level] == [1, 2]


def test_markdown_closing_hashes_are_stripped():
    assert parse_markdown("### Methods ###")[0].text == "Methods"


def test_markdown_setext_headings_are_recognised():
    blocks = parse_markdown("Overview\n========\n\ntext\n\nDetails\n-------\n\nmore\n")

    assert [(block.text, block.level) for block in blocks if block.kind == "heading"] == [
        ("Overview", 1),
        ("Details", 2),
    ]


def test_markdown_code_fences_are_not_parsed_as_headings():
    """
    The guard that matters most here. A Python comment starts with "#", so
    parsing a code sample as Markdown would turn it into a section heading and
    hang the rest of the document underneath it.
    """
    blocks = parse_markdown("Intro.\n\n```python\n# not a heading\nx = 1\n```\n\nOutro.\n")

    assert _kinds(blocks) == ["text", "text", "text"]
    assert "# not a heading" in blocks[1].text


def test_markdown_tables_are_kept_as_tables():
    blocks = parse_markdown("| a | b |\n| --- | --- |\n| 1 | 2 |\n")

    assert _kinds(blocks) == ["table"]


def test_a_lone_pipe_line_is_not_a_table():
    # No divider under it, so it's more likely prose than a header row.
    assert _kinds(parse_markdown("| this is just a line |")) == ["text"]


def test_blank_lines_separate_markdown_paragraphs():
    blocks = parse_markdown("First para\nsecond line.\n\nSecond para.\n")

    assert [block.text for block in blocks] == ["First para\nsecond line.", "Second para."]


def test_plain_text_falls_through_to_paragraphs():
    blocks = parse_markdown("A paragraph.\n\nAnother one.\n")

    assert _kinds(blocks) == ["text", "text"]


# --- HTML -------------------------------------------------------------------


def test_html_heading_tags_carry_their_own_level():
    blocks = parse_html("<h1>One</h1><p>text</p><h3>Three</h3><p>more</p>")

    assert _kinds(blocks) == ["heading", "text", "heading", "text"]
    # h1 and h3 with nothing between them compact to levels 1 and 2 only once
    # rag/structure.py runs; the loader reports what the tags said.
    assert [block.level for block in blocks if block.level] == [1, 3]


def test_script_and_style_contents_are_dropped():
    blocks = parse_html(
        "<head><style>h1 { color: red }</style></head>"
        "<body><script>alert('x')</script><p>Real content.</p></body>"
    )

    assert [block.text for block in blocks] == ["Real content."]


def test_html_tables_become_markdown():
    blocks = parse_html(
        "<table><tr><th>Model</th><th>Acc</th></tr>"
        "<tr><td>BERT</td><td>88.5</td></tr></table>"
    )

    assert _kinds(blocks) == ["table"]
    assert blocks[0].text == "| Model | Acc |\n| --- | --- |\n| BERT | 88.5 |"


def test_ragged_html_table_rows_are_padded():
    # Markdown needs every row to have the same number of cells, and real
    # pages are full of tables where they don't.
    blocks = parse_html("<table><tr><td>a</td><td>b</td></tr><tr><td>c</td></tr></table>")

    assert blocks[0].text.endswith("| c |  |")


def test_source_indentation_does_not_reach_the_text():
    """
    Newlines in HTML source are formatting of the *file*, not the content -
    a paragraph split across indented lines is one line on screen.
    """
    blocks = parse_html("<p>\n    A sentence\n    split across lines.\n</p>")

    assert blocks[0].text == "A sentence split across lines."


def test_entities_are_decoded():
    assert parse_html("<p>Tom &amp; Jerry &lt;3</p>")[0].text == "Tom & Jerry <3"


def test_unclosed_tags_still_yield_their_text():
    # Real pages are full of these, and a strict parser would refuse the file.
    blocks = parse_html("<body><p>First<p>Second")

    assert [block.text for block in blocks] == ["First", "Second"]


# --- Word -------------------------------------------------------------------


def test_docx_heading_styles_become_headings():
    def build(document):
        document.add_heading("Annual Report", level=0)  # the Title style
        document.add_heading("1. Methods", level=1)
        document.add_paragraph("We collected 500 samples.")
        document.add_heading("1.1 Sampling", level=2)
        document.add_paragraph("Sites chosen at random.")

    blocks = load_document("report.docx", _build_docx(build))[0].blocks

    assert _kinds(blocks) == ["heading", "heading", "text", "heading", "text"]
    # Title outranks Heading 1, the same shape font-size ranking gives a PDF.
    assert [block.level for block in blocks if block.level] == [1, 2, 3]


def test_docx_without_a_title_style_still_starts_at_level_one():
    """
    Word's Title maps above Heading 1, so a document with no Title would begin
    at level 2 - and the section stack only closes a section correctly when
    the outermost heading is level 1.
    """
    def build(document):
        document.add_heading("Alpha", level=1)
        document.add_paragraph("a")
        document.add_heading("Beta", level=1)
        document.add_paragraph("b")

    blocks = load_document("report.docx", _build_docx(build))[0].blocks

    assert [block.level for block in blocks if block.level] == [1, 1]
    # Beta replaces Alpha rather than nesting inside it.
    assert blocks[3].section_path == ("Beta",)


def test_docx_tables_keep_their_position_between_paragraphs():
    """
    python-docx exposes paragraphs and tables as two separate lists, which
    loses their interleaving - and a table's position is what says which
    section it belongs to.
    """
    def build(document):
        document.add_heading("Results", level=1)
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Model"
        table.cell(0, 1).text = "Acc"
        table.cell(1, 0).text = "BERT"
        table.cell(1, 1).text = "88.5"
        document.add_paragraph("Closing note.")

    blocks = load_document("report.docx", _build_docx(build))[0].blocks

    assert _kinds(blocks) == ["heading", "table", "text"]
    assert blocks[1].section_path == ("Results",)


def test_empty_docx_paragraphs_are_skipped():
    def build(document):
        document.add_paragraph("Real text.")
        document.add_paragraph("")
        document.add_paragraph("   ")

    assert len(load_document("report.docx", _build_docx(build))[0].blocks) == 1


def test_unreadable_docx_is_rejected():
    with pytest.raises(ValueError, match="Could not read Word document"):
        load_document("broken.docx", b"this is not a docx")


# --- CSV --------------------------------------------------------------------


def test_csv_becomes_a_markdown_table():
    blocks = parse_csv("city,orders\nBergen,4412\nOslo,9931\n")

    assert _kinds(blocks) == ["table"]
    assert blocks[0].text == (
        "| city | orders |\n| --- | --- |\n| Bergen | 4412 |\n| Oslo | 9931 |"
    )


def test_semicolon_delimited_files_are_detected():
    # "CSV" routinely means semicolons in locales where the comma is a decimal
    # separator.
    blocks = parse_csv("city;orders\nBergen;4412\nOslo;9931\n")

    assert blocks[0].text.startswith("| city | orders |")


def test_every_row_group_repeats_the_header():
    """
    Why a CSV isn't emitted as one big table.

    Split at an arbitrary row, a chunk reads "| Bergen | 4412 |" and nothing
    in it says which column is the city and which is the order count. Grouping
    with a repeated header means every chunk carries its own column names.
    """
    rows = "\n".join(f"city{index},{index}" for index in range(ROWS_PER_GROUP * 2 + 3))
    blocks = parse_csv(f"city,orders\n{rows}\n")

    assert len(blocks) == 3
    assert all(block.text.startswith("| city | orders |") for block in blocks)
    assert all(block.kind == "table" for block in blocks)


def test_newlines_inside_quoted_fields_do_not_break_rows():
    blocks = parse_csv('note,value\n"line one\nline two",7\n')

    assert blocks[0].text.endswith("| line one line two | 7 |")


def test_a_header_only_csv_still_ingests():
    # "What columns does this file have" is a reasonable question to ask of it.
    blocks = parse_csv("city,orders,date\n")

    assert _kinds(blocks) == ["table"]
    assert blocks[0].text == "| city | orders | date |\n| --- | --- | --- |"


def test_blank_rows_are_dropped():
    blocks = parse_csv("a,b\n1,2\n\n,\n3,4\n")

    assert blocks[0].text.count("\n") == 3  # header, divider, two data rows
