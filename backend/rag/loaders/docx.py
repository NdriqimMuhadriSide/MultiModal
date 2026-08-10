"""
Word (.docx) loader.

Responsibility: turn a Word document into blocks. Like Markdown and HTML, a
.docx already knows its own structure - a heading is a paragraph whose *style*
is "Heading 2" - so there is nothing to infer from geometry.

A .docx is a zip of XML, so this could be done with the standard library. It
uses python-docx instead, for one reason that isn't laziness: the heading
signal lives in a style *reference*, and resolving it correctly means following
w:pStyle to styles.xml, handling styles that inherit from other styles, and
coping with the fact that the stored name is localised in documents authored in
other languages. python-docx does that resolution properly. Re-implementing it
is a lot of fiddly XML for no insight - unlike, say, rag/layout.py, where the
algorithm *is* the interesting part.

Only the main document body is read: headers, footers, footnotes and comments
are deliberately skipped. A running header repeated on every page is the same
boilerplate rag/layout.py strips out of PDFs, and for the same reason - it
would otherwise turn up inside many chunks and pull them all toward the same
region of embedding space.
"""
from io import BytesIO

# Resolves to the installed python-docx package, not this module: Python 3
# imports are absolute unless written otherwise.
import docx
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from rag.layout import Block, rows_to_markdown
from rag.loaders.base import PageContent, single_page

# Word's built-in outline styles. "Title" is the document title, which sits
# above "Heading 1" in Word's own outline, so it maps to level 1 and pushes
# the numbered chapters below it - the same shape rag/structure.py derives
# from font sizes in a PDF.
_TITLE_STYLES = {"title"}
_HEADING_PREFIX = "heading "


def load_docx(file_bytes: bytes) -> list[PageContent]:
    """
    Parse a .docx into a single page of blocks.

    One page because a .docx has no pages: where they fall depends on the
    printer, the fonts installed, and the zoom level, and Word only computes
    them at render time. There is no page number to cite, so these documents
    are cited by section instead.

    Raises:
        ValueError: if the bytes are not a readable .docx. Callers (the
            service layer) turn this into an HTTP 400.
    """
    try:
        document = docx.Document(BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001 - python-docx raises several types
        raise ValueError(f"Could not read Word document: {exc}") from exc

    return single_page(_blocks(document))


def _blocks(document: Document) -> list[Block]:
    blocks: list[Block] = []
    for item in _body_in_order(document):
        if isinstance(item, Table):
            rows = [[cell.text.strip() for cell in row.cells] for row in item.rows]
            if any(any(cell for cell in row) for row in rows):
                blocks.append(Block(kind="table", text=rows_to_markdown(rows)))
            continue

        text = " ".join(item.text.split())
        if not text:
            continue
        level = _heading_level(item)
        if level:
            blocks.append(Block(kind="heading", text=text, level=level))
        else:
            blocks.append(Block(kind="text", text=text))
    return blocks


def _body_in_order(document: Document):
    """
    Yield the body's paragraphs and tables in the order they appear.

    python-docx exposes `document.paragraphs` and `document.tables` as two
    separate lists, which loses their interleaving - and a table's position is
    exactly what tells you which section it belongs to. Walking the underlying
    XML children is the only way to keep document order.
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def _heading_level(paragraph: Paragraph) -> int | None:
    """
    Read a paragraph's outline level from its style name.

    Word's Title style sits *above* Heading 1 in its own outline, so it takes
    level 1 and the numbered headings shift down one. A document with no Title
    therefore starts at 2 - which is fine, because rag/structure.py compacts
    the levels a document actually uses back down to start at 1.

    Falls back to the paragraph's recorded outline level where there is one: a
    custom style can be a heading without being called "Heading 3".
    """
    name = (paragraph.style.name or "").strip().lower() if paragraph.style else ""

    if name in _TITLE_STYLES:
        return 1
    if name.startswith(_HEADING_PREFIX):
        depth = name[len(_HEADING_PREFIX) :].strip()
        if depth.isdigit():
            return int(depth) + 1

    outline = _outline_level(paragraph)
    # Word stores outline levels 0-8 (0 being Heading 1), where 9 means
    # "body text" - i.e. not a heading at all.
    return outline + 2 if outline is not None and outline < 9 else None


def _outline_level(paragraph: Paragraph) -> int | None:
    try:
        outline = paragraph.style.paragraph_format.element.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl"
        )
    except Exception:  # noqa: BLE001 - styles can be missing or malformed
        return None
    if outline is None:
        return None
    value = outline.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
    return int(value) if value is not None and value.isdigit() else None
